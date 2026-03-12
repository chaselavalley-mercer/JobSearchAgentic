import sys
import json
import os
import re
from docx import Document  
from docx.enum.text import WD_TAB_ALIGNMENT

def delete_paragraph(paragraph):
    """
    Deletes a paragraph from the document completely to ensure no leftover whitespace.
    """
    p = paragraph._element
    if p.getparent() is not None:
        p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def merge_runs(paragraph):
    """
    Merges adjacent runs that have identical character formatting.
    This defragments placeholders that Word might have split across multiple runs.
    """
    if len(paragraph.runs) <= 1:
        return

    def get_run_style_key(run):
        try:
            return (
                run.style.name if run.style else None,
                run.bold,
                run.italic,
                run.underline,
                run.font.name if run.font else None,
                run.font.size if run.font else None,
                run.font.color.rgb if (run.font and run.font.color and run.font.color.type == 1) else None
            )
        except Exception:
            return None

    current_run = paragraph.runs[0]
    current_key = get_run_style_key(current_run)
    
    for i in range(1, len(paragraph.runs)):
        run = paragraph.runs[i]
        key = get_run_style_key(run)
        
        if key == current_key and key is not None:
            # Merge text into the current_run and clear this run's text
            if run.text:
                current_run.text += run.text
                run.text = ""
        else:
            # Start a new sequence
            current_key = key
            current_run = run

def regex_replace_in_paragraph(paragraph, pattern_str, value):
    """
    Finds regex pattern in paragraph.text and replaces the first occurrence 
    with value, respecting run boundaries.
    Single-pass only — called once per slot per paragraph, as one replacement 
    is sufficient and prevents infinite loops.
    """
    full_text_chars = []
    char_map = []
    for r_idx, run in enumerate(paragraph.runs):
        run_text = getattr(run, "text", "")
        if isinstance(run_text, str):
            for c_idx in range(len(run_text)):
                char_map.append((r_idx, c_idx))
                full_text_chars.append(run_text[c_idx])
            
    full_text = "".join(full_text_chars)
    match = re.search(pattern_str, full_text)
    if not match:
        return
        
    start_idx = match.start()
    end_idx = match.end() - 1
    
    if start_idx > end_idx:
        # Empty match string, prevent infinite loop safely
        return
        
    start_run_idx = char_map[start_idx][0]
    end_run_idx = char_map[end_idx][0]
    
    start_char_idx = char_map[start_idx][1]
    end_char_idx = char_map[end_idx][1]
    
    if start_run_idx == end_run_idx:
        run = paragraph.runs[start_run_idx]
        prefix = run.text[:start_char_idx]
        suffix = run.text[end_char_idx+1:]
        run.text = prefix + value + suffix
    else:
        start_run = paragraph.runs[start_run_idx]
        end_run = paragraph.runs[end_run_idx]
        
        start_run_prefix = start_run.text[:start_char_idx]
        end_run_suffix = end_run.text[end_char_idx+1:]
        
        start_run.text = start_run_prefix + value
        end_run.text = end_run_suffix
        
        # Clear intermediate runs
        for i in range(start_run_idx + 1, end_run_idx):
            paragraph.runs[i].text = ""

def process_paragraph(paragraph, mapping, usable_width=None):
    """
    Returns True if paragraph should be marked for deletion.
    Handles searching full text, replacing slots, and cleaning typography dynamically.
    """
    merge_runs(paragraph)
    full_text = paragraph.text
    matches = list(set(re.findall(r'\[[A-Za-z][A-Za-z0-9_]+\]', full_text)))
    
    if not matches:
        return False
        
    print(f"Scanning Paragraph: {full_text[:60]}{'...' if len(full_text)>60 else ''}")
    print(f"  Found Slots: {matches}")
        
    all_empty = True
    slot_evals = {}
    
    for match in matches:
        # Resolve value (checks mapped string with AND without brackets)
        val = mapping.get(match)
        if val is None:
            core_key = match[1:-1]
            val = mapping.get(core_key)
            
        is_empty = (val is None or str(val).strip() == "")
        sanitized_val = "" if is_empty else str(val)
        
        # Sanitize specific fields to ensure they stay strictly on one line
        if any(k in match.upper() for k in ['DATE', 'EMPLOYER', 'LOCATION', 'COMPANY']):
            sanitized_val = re.sub(r'\s+', ' ', sanitized_val).strip()
            
        slot_evals[match] = {
            "empty": is_empty,
            "value": sanitized_val
        }
        
        if not is_empty:
            all_empty = False
            
    if all_empty:
        print("  Action: Delete")
        return True
        
    print("  Action: Keep")
    
    # Date alignment styling intelligence
    date_slots = [m for m in matches if 'Date' in m or 'DATE' in m]
    loc_emp_slots = [m for m in matches if any(k in m.upper() for k in ['LOCATION', 'EMPLOYER', 'COMPANY'])]
    
    if date_slots and loc_emp_slots:
        for d_slot in date_slots:
            escaped_d = r'\[' + re.escape(d_slot[1:-1]) + r'\]'
            # Clean ALL whitespace (including \n, \r) before date slot, replacing with a single tab
            regex_replace_in_paragraph(paragraph, r'\s+' + escaped_d, '\t' + d_slot)
            
        has_right_tab = False
        try:
            if paragraph.paragraph_format and paragraph.paragraph_format.tab_stops:
                for ts in paragraph.paragraph_format.tab_stops:
                    if getattr(ts, 'alignment', None) == WD_TAB_ALIGNMENT.RIGHT:
                        has_right_tab = True
                        break
        except Exception:
            pass
            
        if not has_right_tab and usable_width:
            try:
                # Programmatically enforce a right-aligned tab stop at the true right margin
                paragraph.paragraph_format.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
            except Exception as e:
                print(f"  [!] Could not add dynamic right tab stop: {e}")
    # Process replacements for this kept paragraph
    for match in matches:
        eval_data = slot_evals[match]
        escaped_slot = r'\[' + re.escape(match[1:-1]) + r'\]'
        
        if eval_data["empty"]:
            # Clean up empty slots and optionally a preceding or succeeding separator
            sep = r'(?:\||-|•|–|—|,)'
            # We prioritize deleting a preceding separator, or a succeeding separator, or just the slot itself.
            pattern = rf'(?:\s*{sep}\s*{escaped_slot})|(?:{escaped_slot}\s*{sep}\s*)|(?:{escaped_slot})'
            regex_replace_in_paragraph(paragraph, pattern, "")
        else:
            # Populate valid slot data
            regex_replace_in_paragraph(paragraph, escaped_slot, eval_data["value"])
            
    if not paragraph.text.strip():
        print("  Action: Delete (Became empty after replacements)")
        return True
            
    return False

def main():
    if len(sys.argv) != 3:
        print("Usage: python remodel_docx.py <output_file> <json_file_path>")
        sys.exit(1)

    output_file = sys.argv[1]
    json_file_path = sys.argv[2]
    
    # Deterministic resolution of template path independent of CWD
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    template_path = os.path.join(base_dir, "directives", "global_master.docx")

    if not os.path.exists(template_path):
        print(f"Error: Template file not found at {template_path}")
        sys.exit(1)

    if not os.path.exists(json_file_path):
        print(f"Error: JSON file not found at {json_file_path}")
        sys.exit(1)

    try:
        with open(json_file_path, 'r', encoding='utf-8') as f:
            mapping = json.load(f)
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON format in file. Details: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Error reading JSON file: {e}")
        sys.exit(1)

    if not isinstance(mapping, dict):
        print("Error: JSON mapping must be a dictionary object.")
        sys.exit(1)
        
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"Error loading template docx: {e}")
        sys.exit(1)
        
    # Calculate usable width dynamically from the fully verified doc object once
    try:
        section = doc.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
    except Exception:
        usable_width = None

    paragraphs_to_delete = []

    # Process all regular paragraphs
    for p in doc.paragraphs:
        if process_paragraph(p, mapping, usable_width):
            paragraphs_to_delete.append(p)
            
    # Process all table paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if process_paragraph(p, mapping, usable_width):
                        paragraphs_to_delete.append(p)

    # Clean up empty paragraphs safely
    for p in paragraphs_to_delete:
        delete_paragraph(p)
        
    # Standardize output directory
    output_abs_path = os.path.abspath(output_file)
    os.makedirs(os.path.dirname(output_abs_path), exist_ok=True)

    try:
        doc.save(output_abs_path)
        print(f"Success: Remodeled document saved to {output_abs_path}")
    except Exception as e:
        print(f"Error saving document to {output_abs_path}: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
