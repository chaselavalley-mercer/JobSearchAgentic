"""
docx_assembler.py — Phase 3 Resume Assembler
Layer: 3 (Execution)

Usage:
  python execution/docx_assembler.py \\
    --template directives/global_master.docx \\
    --mapping  .tmp/{user_id}/resume_args.json \\
    --output   .users/{user_id}/{user_id}_{job_title}_{company}.docx

Exit codes:
  0  — Document assembled within page limit (~1.5 pages).
  1  — Fatal error (file missing, bad JSON, subprocess failure).
  2  — TRIM_REQUIRED: Assembled document exceeds page limit.
       Orchestrator must drop lowest-scored project and re-run.

Design:
  This script is a smart wrapper around execution/remodel_docx.py.
  It handles argument parsing, env loading, subprocess invocation,
  and page-count estimation. Low-level paragraph surgery stays in
  remodel_docx.py — the two scripts remain independently testable.
"""

import sys
import os
import json
import argparse
import subprocess

# ---------------------------------------------------------------------------
# Self-Source: load .env from project root
# ---------------------------------------------------------------------------
try:
    from dotenv import load_dotenv
    _env_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), ".env")
    load_dotenv(dotenv_path=_env_path, override=False)
except ImportError:
    pass

# ---------------------------------------------------------------------------
# Configuration constants
# ---------------------------------------------------------------------------
PYTHON_PATH = os.environ.get(
    "PYTHON_PATH",
    r"C:\Users\cal-asus1\AppData\Local\Programs\Python\Python311\python.exe"
)

# Words-per-page heuristic for a dense resume (10pt, 0.5in margins).
# Calibrated for single-spaced, bullet-heavy resumes.
WORDS_PER_PAGE = 450

# Hard page ceiling. Exceeding this triggers TRIM_REQUIRED (exit 2).
PAGE_CEILING = 1.5

# Script that performs actual docx paragraph surgery
REMODEL_SCRIPT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "remodel_docx.py")


def estimate_page_count(docx_path: str) -> float:
    """
    Estimate page count of a .docx file using a word-count heuristic.
    Returns a float (e.g. 1.3). Not a substitute for Word's renderer,
    but sufficient for the Trim Loop gate.
    """
    try:
        from docx import Document
        doc = Document(docx_path)
        total_words = sum(
            len(para.text.split())
            for para in doc.paragraphs
            if para.text.strip()
        )
        # Also count table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        total_words += len(para.text.split())
        estimate = total_words / WORDS_PER_PAGE
        return estimate
    except Exception as e:
        print(f"[WARN] Could not estimate page count: {e}. Defaulting to 1.0.")
        return 1.0


def main():
    parser = argparse.ArgumentParser(description="Phase 3 Resume Assembler")
    parser.add_argument("--template", required=True, help="Path to global_master.docx template")
    parser.add_argument("--mapping", required=True, help="Path to resume_args.json slot mapping")
    parser.add_argument("--output", required=True, help="Path for the output .docx file")
    args = parser.parse_args()

    # ------------------------------------------------------------------
    # Pre-flight checks
    # ------------------------------------------------------------------
    missing = []
    if not os.path.exists(args.template):
        missing.append(f"Template not found: {args.template}")
    if not os.path.exists(args.mapping):
        missing.append(f"Mapping file not found: {args.mapping}")
    if not os.path.exists(REMODEL_SCRIPT):
        missing.append(f"remodel_docx.py not found: {REMODEL_SCRIPT}")
    if not os.path.exists(PYTHON_PATH):
        missing.append(
            f"Python interpreter not found: {PYTHON_PATH}. "
            "Set PYTHON_PATH in .env or update the PYTHON_PATH constant."
        )

    if missing:
        for m in missing:
            print(f"[ERROR] {m}")
        sys.exit(1)

    # Validate mapping JSON is parseable before handing to subprocess
    try:
        with open(args.mapping, "r", encoding="utf-8") as f:
            mapping_data = json.load(f)
        if not isinstance(mapping_data, dict):
            print("[ERROR] resume_args.json must be a JSON object (dict).")
            sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"[ERROR] Invalid JSON in mapping file: {e}")
        sys.exit(1)

    # ------------------------------------------------------------------
    # Ensure output directory exists
    # ------------------------------------------------------------------
    output_abs = os.path.abspath(args.output)
    os.makedirs(os.path.dirname(output_abs), exist_ok=True)

    # ------------------------------------------------------------------
    # Invoke remodel_docx.py as subprocess
    # The template path is baked into remodel_docx.py's resolution logic
    # (it resolves relative to its own directory), so we pass output + mapping.
    # ------------------------------------------------------------------
    print(f"[INFO] Assembling resume...")
    print(f"  Template : {os.path.abspath(args.template)}")
    print(f"  Mapping  : {os.path.abspath(args.mapping)}")
    print(f"  Output   : {output_abs}")

    cmd = [PYTHON_PATH, REMODEL_SCRIPT, output_abs, os.path.abspath(args.mapping)]

    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace"
        )
    except Exception as e:
        print(f"[ERROR] Failed to launch remodel_docx.py subprocess: {e}")
        sys.exit(1)

    # Relay stdout/stderr from subprocess
    if result.stdout:
        print(result.stdout.rstrip())
    if result.stderr:
        print(result.stderr.rstrip(), file=sys.stderr)

    if result.returncode != 0:
        print(f"[ERROR] remodel_docx.py exited with code {result.returncode}.")
        sys.exit(1)

    if not os.path.exists(output_abs):
        print(f"[ERROR] Expected output file was not created: {output_abs}")
        sys.exit(1)

    # ------------------------------------------------------------------
    # Page count gate
    # ------------------------------------------------------------------
    estimated_pages = estimate_page_count(output_abs)
    print(f"[INFO] Estimated page count: {estimated_pages:.2f} (ceiling: {PAGE_CEILING})")

    if estimated_pages > PAGE_CEILING:
        print(
            f"[TRIM_REQUIRED] Document exceeds {PAGE_CEILING} page limit "
            f"({estimated_pages:.2f} estimated pages).\n"
            "  Orchestrator: drop the lowest-ranked project slots and re-run."
        )
        sys.exit(2)

    print(f"[SUCCESS] Resume assembled at: {output_abs}")
    sys.exit(0)


if __name__ == "__main__":
    main()
