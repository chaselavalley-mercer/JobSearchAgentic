"""
audit_resume.py — Phase 4 Hallucination Check & Pipeline Logger
Layer: 3 (Execution)

Usage:
  python execution/audit_resume.py \\
    --docx     .users/{user_id}/{filename}.docx \\
    --knowledge .users/{user_id}/knowledge.md \\
    --log-dir  logs/{user_id}/

Exit codes:
  0  — Audit PASSED. Log written to log-dir.
  1  — Audit FAILED. Flagged bullets printed and written to log. Fix and re-run.
  2  — Runtime error (missing file, bad docx, etc.).

Hallucination Detection Strategy:
  The script extracts all non-empty text lines from the .docx and checks
  each against a reference corpus built from knowledge.md. A line is flagged
  if it contains a "claim token" (company name, proper noun, metric, or
  technology keyword) that cannot be found anywhere in the knowledge corpus.

  This is intentionally conservative — false negatives (missed hallucinations)
  are worse than false positives (flagged-but-valid lines). The human auditor
  makes the final call.
"""

import sys
import os
import re
import json
import argparse
from datetime import datetime

# ---------------------------------------------------------------------------
# Self-Source: load .env from project root
# ---------------------------------------------------------------------------
try:
    from dotenv import load_dotenv
    _env_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), ".env")
    load_dotenv(dotenv_path=_env_path, override=False)
    _enc = os.environ.get("PYTHONIOENCODING", "")
    if _enc:
        sys.stdout.reconfigure(encoding=_enc)
        sys.stderr.reconfigure(encoding=_enc)
except (ImportError, AttributeError):
    pass

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
PYTHON_PATH = os.environ.get(
    "PYTHON_PATH",
    r"C:\Users\cal-asus1\AppData\Local\Programs\Python\Python311\python.exe"
)

# Minimum token length for entity tokens
MIN_CLAIM_TOKEN_LEN = 3

# Tokens to always ignore when building the knowledge corpus
STOPWORDS = {
    "with", "using", "and", "the", "for", "from", "that", "this", "have",
    "been", "were", "will", "into", "over", "under", "their", "which", "while",
    "built", "used", "made", "work", "team", "role", "skills", "time", "data",
    "systems", "based", "across", "within", "through", "improved", "managed",
    "developed", "designed", "implemented", "created", "delivered", "supported",
    "university", "experience", "engineering", "software", "hardware", "project",
    "projects", "technical", "including", "results", "during", "strong", "key",
    "lead", "led", "build", "test", "ensure", "provide", "drive"
}

# Always valid — common tech abbreviations and resume vocabulary
KNOWN_SAFE = {
    "api", "apis", "rest", "http", "https", "json", "sql", "cli",
    "ide", "sdk", "csv", "pdf", "html", "css", "js", "url", "ai",
    "ml", "llm", "cv", "gpa", "bs", "bse", "ga", "llc", "inc",
    "github", "linkedin", "www", "com", "edu", "live"
}

# Excluded from entity extraction even though they start with capitals
EXCLUDE_FROM_ENTITIES = {
    # Months
    "january", "february", "march", "april", "may", "june",
    "july", "august", "september", "october", "november", "december",
    "jan", "feb", "mar", "apr", "jun", "jul", "aug", "sep", "oct", "nov", "dec",
    # Days
    "monday", "tuesday", "wednesday", "thursday", "friday", "saturday", "sunday",
    "mon", "tue", "wed", "thu", "fri", "sat", "sun",
    # Common suffixes/titles
    "inc", "llc", "ltd", "corp",
}

# Common resume action verbs (sentence-starting) and generic title-case nouns
# that should never be treated as entity claims regardless of capitalization.
# Stored lowercase; compared via token.lower().
COMMON_ENGLISH_CAPS = {
    # Action verbs that start resume bullets
    "designed", "built", "engineered", "developed", "implemented", "deployed",
    "created", "architected", "managed", "led", "coached", "taught", "mentored",
    "planned", "executed", "integrated", "optimized", "reduced", "increased",
    "delivered", "demonstrated", "achieved", "ensured", "coordinated",
    "standardized", "clarified", "instructed", "reinforced", "fusing",
    "producing", "achieving", "enabling", "managing",
    # Generic nouns appearing in skill headers and bullet prose
    "vision", "pipelines", "pipeline", "integration", "services", "outputs",
    "inputs", "platforms", "analysis", "generation", "architecture", "execution",
    "processing", "production", "systems", "data", "tools", "skills", "labs",
    "reports", "operations", "performance", "events", "tracking", "control",
    "algorithm", "framework", "protocol", "interface", "platform", "database",
    "storage", "workflow", "sequence", "logic", "documentation", "development",
    "fundamentals", "procedures", "checklists", "guidance", "insights",
    "behavior", "circuit", "sessions", "stores",
    "engineering", "university",
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def extract_docx_lines(docx_path: str) -> list[str]:
    """Extract all non-empty text lines from a .docx file."""
    try:
        from docx import Document
    except ImportError:
        print("[ERROR] python-docx is required. Install: pip install python-docx")
        sys.exit(2)

    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        lines.append(text)
    return lines


def build_knowledge_corpus(knowledge_path: str) -> set[str]:
    """
    Build a lowercase set of all meaningful tokens from knowledge.md.
    This is the ground truth against which the resume is checked.
    """
    with open(knowledge_path, "r", encoding="utf-8") as f:
        raw = f.read()

    # Tokenize: split on non-alphanumeric, lowercase, filter short/stopword tokens
    # Refined regex to avoid capturing trailing punctuation unless it's part of a tech tag
    tokens = re.findall(r"[A-Za-z0-9][A-Za-z0-9\+\#\.]*", raw)
    corpus = set()
    for tok in tokens:
        # Strip trailing periods if they aren't part of a known tech pattern like .NET
        processed_tok = tok.rstrip('.')
        if not processed_tok:
            continue
        lower = processed_tok.lower()
        if len(lower) >= MIN_CLAIM_TOKEN_LEN and lower not in STOPWORDS:
            corpus.add(lower)

    return corpus


def extract_claim_tokens(line: str) -> list[str]:
    """
    Extract entity tokens from a resume bullet line.
    Only returns tokens that look like entities — proper nouns, abbreviations,
    metrics, or technology names. Common English words are never returned.

    Rules:
    - Must be 3+ characters
    - Must start with a capital letter, be all-caps, or contain digits
    - Skipped if in KNOWN_SAFE or EXCLUDE_FROM_ENTITIES
    """
    tokens = re.findall(r"[A-Za-z0-9][A-Za-z0-9\+\#\.]*", line)
    claims = []
    for tok in tokens:
        processed = tok.rstrip('.')
        if not processed or len(processed) < MIN_CLAIM_TOKEN_LEN:
            continue
        lower = processed.lower()
        if lower in KNOWN_SAFE or lower in EXCLUDE_FROM_ENTITIES:
            continue
        # Entity filter: capitals, all-caps, or contains digits
        is_capitalized = processed[0].isupper()
        is_allcaps = processed.isupper() and len(processed) >= 2
        has_digits = any(c.isdigit() for c in processed)
        if not (is_capitalized or is_allcaps or has_digits):
            continue
        # Skip common English words that appear capitalized but aren't entities
        if lower in COMMON_ENGLISH_CAPS:
            continue
        claims.append(lower)
    return claims


def audit_lines(
    resume_lines: list[str],
    corpus: set[str]
) -> tuple[list[str], list[dict], int]:
    """
    Compare resume lines against the knowledge corpus using entity-token matching.
    Returns:
      - clean_lines: lines that fully pass
      - flagged_lines: list of {line, unknown_tokens} for lines with unverified entity claims
      - total_entities: total count of entity tokens checked across all lines
    """
    clean = []
    flagged = []
    total_entities = 0

    for line in resume_lines:
        claims = extract_claim_tokens(line)
        total_entities += len(claims)
        unknown = [tok for tok in claims if tok not in corpus]

        # Only flag substantive lines (6+ words) with unknown entity tokens
        if unknown and len(line.split()) >= 6:
            flagged.append({"line": line, "unknown_tokens": unknown})
        else:
            clean.append(line)

    return clean, flagged, total_entities


def write_log(
    log_dir: str,
    docx_path: str,
    knowledge_path: str,
    flagged: list[dict],
    resume_lines: list[str],
    passed: bool,
    entity_stats: dict
) -> str:
    """Write a structured audit log and return the log file path."""
    os.makedirs(log_dir, exist_ok=True)

    company_guess = os.path.basename(docx_path).replace(".docx", "").replace("_", " ")
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    log_filename = f"{datetime.now().strftime('%Y-%m-%d')}_{company_guess}.log"
    log_path = os.path.join(log_dir, log_filename)

    status = "PASS" if passed else "FAIL"

    lines_out = [
        "=" * 72,
        f"AUDIT LOG — {status}",
        "=" * 72,
        f"Timestamp  : {timestamp}",
        f"DOCX       : {os.path.abspath(docx_path)}",
        f"Knowledge  : {os.path.abspath(knowledge_path)}",
        f"Status     : {status}",
        f"Lines      : {len(resume_lines)} total resume lines audited",
        f"Flagged    : {len(flagged)} lines with unverified claim tokens",
        f"Entity tokens checked : {entity_stats['checked']} | Flagged: {entity_stats['flagged']}",
        "",
    ]

    if flagged:
        lines_out.append("FLAGGED LINES (review for hallucinations):")
        lines_out.append("-" * 72)
        for i, item in enumerate(flagged, 1):
            lines_out.append(f"  [{i}] {item['line']}")
            lines_out.append(f"      Unknown tokens: {', '.join(item['unknown_tokens'])}")
        lines_out.append("")

    lines_out += [
        "FINAL STATE SNAPSHOT",
        "-" * 72,
        f"Total clean lines  : {len(resume_lines) - len(flagged)}",
        f"Total flagged lines: {len(flagged)}",
        f"Pipeline result    : {'Resume is ready.' if passed else 'Fix flagged lines and re-run.'}",
        "=" * 72,
        "",
    ]

    with open(log_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines_out))

    return log_path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Phase 4 Hallucination Check & Logger")
    parser.add_argument("--docx", required=True, help="Path to the assembled output .docx")
    parser.add_argument("--knowledge", required=True, help="Path to user's knowledge.md")
    parser.add_argument("--log-dir", required=True, help="Directory to write the audit log")
    args = parser.parse_args()

    # Pre-flight
    for path, label in [(args.docx, "DOCX"), (args.knowledge, "knowledge.md")]:
        if not os.path.exists(path):
            print(f"[ERROR] {label} not found: {path}")
            sys.exit(2)

    print(f"[INFO] Running audit...")
    print(f"  DOCX      : {args.docx}")
    print(f"  Knowledge : {args.knowledge}")
    print(f"  Log dir   : {args.log_dir}")

    resume_lines = extract_docx_lines(args.docx)
    corpus = build_knowledge_corpus(args.knowledge)

    print(f"[INFO] Auditing {len(resume_lines)} lines against "
          f"{len(corpus)} knowledge tokens...")

    _, flagged, total_entities = audit_lines(resume_lines, corpus)
    passed = len(flagged) == 0

    entity_stats = {"checked": total_entities, "flagged": sum(len(f["unknown_tokens"]) for f in flagged)}

    log_path = write_log(
        log_dir=args.log_dir,
        docx_path=args.docx,
        knowledge_path=args.knowledge,
        flagged=flagged,
        resume_lines=resume_lines,
        passed=passed,
        entity_stats=entity_stats
    )

    if passed:
        print(f"[AUDIT PASS] No hallucinations detected.")
        print(f"[LOG] Written to: {log_path}")
        sys.exit(0)
    else:
        print(f"[AUDIT FAIL] {len(flagged)} line(s) flagged for review:")
        for i, item in enumerate(flagged, 1):
            print(f"  [{i}] {item['line'][:80]}{'...' if len(item['line']) > 80 else ''}")
            print(f"       Unknown: {', '.join(item['unknown_tokens'])}")
        print(f"[LOG] Written to: {log_path}")
        print("  Fix flagged bullets in resume_args.json and re-run Phase 3 + Phase 4.")
        sys.exit(1)


if __name__ == "__main__":
    main()
